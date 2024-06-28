import os
import re
import xlrd
import easyocr
from PIL import Image
from docx import Document
from docx2pdf import convert
from gtts import gTTS
from openpyxl import load_workbook
from pdfminer.high_level import extract_text

def sanitize_xml_string(text):
    sanitized = text.replace('\x00', '')
    sanitized = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', sanitized)
    return sanitized

def pdf_to_word(input_file, output_file):
    try:
        text = extract_text(input_file)
        sanitized_text = sanitize_xml_string(text)
        doc = Document()
        doc.add_paragraph(sanitized_text)
        doc.save(output_file)
    except Exception as e:
        print(f"Error converting PDF to Word: {e}")

def word_to_pdf(input_file):
    output_file = os.path.join('output', 'converted_' + os.path.splitext(os.path.basename(input_file))[0] + '.pdf')
    convert(input_file, output_file)
    return output_file

def xl_to_word(input_file, output_file):
    try:
        if input_file.endswith('.xls'):
            wb = xlrd.open_workbook(input_file)
            ws = wb.sheet_by_index(0)
            doc = Document()
            for row in range(ws.nrows):
                row_text = ' '.join([str(ws.cell_value(row, col)) for col in range(ws.ncols)])
                doc.add_paragraph(row_text)
        else:
            wb = load_workbook(input_file)
            ws = wb.active
            doc = Document()
            for row in ws.iter_rows():
                row_text = ' '.join([str(cell.value) for cell in row])
                doc.add_paragraph(row_text)
        doc.save(output_file)
    except Exception as e:
        print(f"Error converting Excel to Word: {e}")

def image_to_pdf(input_file):
    output_file = os.path.join('output', 'converted_' + os.path.splitext(os.path.basename(input_file))[0] + '.pdf')
    image = Image.open(input_file)
    image.save(output_file, 'PDF')
    return output_file

def easyocr_image_text_extractor(image_path):
    reader = easyocr.Reader(['en'])  
    result = reader.readtext(image_path)
    extracted_text = ' '.join([item[1] for item in result])
    return extracted_text

def image_text_extractor(input_file, output_file):
    try:
        text = easyocr_image_text_extractor(input_file)
        with open(output_file, 'w') as f:
            f.write(text)
    except Exception as e:
        print(f"Error extracting text from Image: {e}")



